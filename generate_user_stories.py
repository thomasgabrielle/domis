from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Jamaica School Feeding eMIS Module\nUser Stories')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Derived from Business Requirements Document v1.0\n\nSDG Joint Programme on Digital Transformation for Education\nWorld Food Programme / Ministry of Education, Youth and Information\nMarch 2026')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_page_break()

def story(sid, title, role, want, so_that, acs, trace):
    doc.add_heading(f'{sid}: {title}', level=3)
    p = doc.add_paragraph()
    p.add_run('As a ')
    r = p.add_run(role)
    r.bold = True
    p.add_run(', I want ')
    p.add_run(want)
    p.add_run(', so that ')
    p.add_run(so_that + '.')
    ac = doc.add_paragraph()
    r = ac.add_run('Acceptance Criteria:')
    r.bold = True
    for a in acs:
        doc.add_paragraph(a, style='List Bullet')
    p = doc.add_paragraph()
    r = p.add_run('BRD Traceability: ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r = p.add_run(trace)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ===== 1. School Profile =====
doc.add_heading('1. School Profile \u2014 School Feeding Data', level=1)

story('US-SP-01', 'Configure School Feeding profile',
    'School Administrator',
    'select School Feeding Service Types (School Meals, Education, Garden, Other) for my school',
    'the SF Unit knows which services my school participates in',
    ['Multi-select enumeration field', 'Updated annually', 'Only School Administrators and above can edit'],
    'SP-01')

story('US-SP-02', 'Designate SF Coordinator(s)',
    'School Administrator',
    'designate one or more School Feeding Coordinators from personnel records (name and title/role)',
    'the SF Unit knows who manages the programme at my school',
    ['Selection from existing personnel records', 'Supports multiple coordinators', 'Updated annually'],
    'SP-02')

story('US-SP-03', 'Record School Nurse',
    'School Administrator',
    'record the School Nurse name from personnel records',
    'the health contact for the school is on file',
    ['Optional field', 'Selection from personnel records', 'Updated annually'],
    'SP-03')

story('US-SP-04', 'Record NPL status',
    'School Administrator',
    'indicate whether my school is an NPL (Nutrition Products Limited) school',
    'the SF Unit can track NPL coverage',
    ['Boolean yes/no', 'Updated annually'],
    'SP-04')

story('US-SP-05', 'Configure meal types',
    'School Administrator',
    'select which meal types my school offers (breakfast, snack, lunch)',
    'the SF Unit knows what meals are being provided',
    ['Multi-select enumeration', 'Updated termly'],
    'SP-05')

story('US-SP-06', 'Record school garden status',
    'School Administrator',
    'indicate whether a school garden is established',
    'the SF Unit can track garden programme participation',
    ['Boolean yes/no', 'Updated annually'],
    'SP-06')

story('US-SP-07', 'Record meal sources',
    'School Administrator',
    'select meal sources (hypermarket, local market, farmer, service provider, school garden, other)',
    'the SF Unit can monitor food sourcing patterns',
    ['Multi-select enumeration', 'Updated termly'],
    'SP-07')

doc.add_page_break()

# ===== 2. Student Registry =====
doc.add_heading('2. Student Registry \u2014 PATH, Welfare, and Wards of the State', level=1)

story('US-SR-01', 'Assign SF programme type to student',
    'SF School Coordinator',
    'assign a School Feeding Programme type (PATH, Wards of the State, Welfare, Other) to each student',
    'every student eligibility category is recorded',
    ['Single-select enumeration per student', 'Editable by SF Coordinator and above'],
    'SR-01')

story('US-SR-02', 'Capture PATH Household ID',
    'SF School Coordinator',
    'record the PATH Household ID for PATH students under their family/guardian information',
    'the student can be linked to the MLSS PATH programme',
    ['Alphanumeric field', 'Stored under student family/guardian data', 'PATH student IDs must be unique; duplicates flagged (VR-05)'],
    'SR-02, VR-05')

story('US-SR-03', 'Capture PATH benefit dates',
    'SF School Coordinator',
    'record PATH benefit start and end dates per student',
    'the system knows when a student is actively receiving PATH benefits',
    ['Date range fields (yyyy-mm-dd)', 'Start date must be before end date (VR-06)'],
    'SR-03, VR-06')

story('US-SR-04', 'Bulk import PATH beneficiary lists',
    'MLSS Regional Social Worker',
    'upload PATH beneficiary lists via CSV/Excel',
    'student programme attributions can be updated in batch',
    ['CSV and Excel formats supported', 'Validation checks on import (invalid IDs, mismatched school codes)', 'Error reporting with details of rejected records', 'All imports logged with timestamp, source, and user'],
    'SR-04, Section 9.3')

story('US-SR-05', 'Flag PATH list discrepancies',
    'SF School Coordinator',
    'the system to flag discrepancies between the imported PATH list and school enrollment',
    'I can investigate and resolve mismatches',
    ['Flag students on PATH list not enrolled at school', 'Flag enrolled students not on PATH list', 'Discrepancy list viewable after each import'],
    'SR-05')

story('US-SR-06', 'View programme attribution history',
    'SF Unit Data Manager',
    'the system to maintain a historical log of changes to programme attribution with date and user',
    'I can audit eligibility changes',
    ['Log: change type, old value, new value, date, user', 'Accessible to HQ-level users'],
    'SR-06')

story('US-SR-07', 'Search and filter student registry',
    'SF Unit Director',
    'filter and search the student registry by programme type, school, region, and parish',
    'I can find specific student populations quickly',
    ['Filter by programme type', 'Filter by school, region, parish', 'Search by student name or ID'],
    'SR-07')

story('US-SR-08', 'Include Welfare and Wards in registry',
    'SF School Coordinator',
    'register Welfare students and Wards of the State in the SF registry',
    'they can be cross-referenced with MLSS for PATH eligibility verification',
    ['Same registration workflow as PATH students', 'Programme type clearly distinguished'],
    'SR-08')

doc.add_page_break()

# ===== 3. Attendance =====
doc.add_heading('3. Student Attendance and PATH Compliance', level=1)

story('US-AT-01', 'Enter daily student attendance',
    'School Administrator / SF Coordinator',
    'enter daily student attendance (or draw from existing eMIS attendance data)',
    'attendance is captured without duplicate entry',
    ['Class-by-class or whole-school entry', 'Bulk mark support (mark all present)', 'Responsive for up to 2,000 students (NF-03)', 'Must work offline with auto-sync (NF-09)'],
    'AT-01, NF-03, NF-09')

story('US-AT-02', 'Calculate attendance rates',
    'SF School Coordinator',
    'the system to calculate student attendance rates monthly and bi-monthly',
    'compliance can be monitored regularly',
    ['Formula: (days attended / total school days) x 100 (VR-08)', 'Auto-calculated from attendance records'],
    'AT-02, VR-08')

story('US-AT-03', 'Auto-flag non-compliant PATH students',
    'MLSS Director',
    'the system to automatically identify PATH students below 85% attendance',
    'benefit adjustments can be applied',
    ['PATH students aged 6-18 (BR-01)', 'Vacation months Jul-Aug = auto-compliant (BR-02)', 'Illness code 6 not penalized (BR-03)', 'Chronic illness = special case with medical report (BR-04)', 'Disability in special institutions = exempt (BR-05)', 'Violence/disaster absences = compliant with written notice (BR-06)'],
    'AT-03, BR-01 to BR-06')

story('US-AT-04', 'Generate Attendance Verification Record',
    'MLSS Associate Officer',
    'the system to generate an Attendance Verification Record per school bi-monthly',
    'MLSS receives compliance data without manual submission',
    ['Auto-generated bi-monthly', 'Extractable as report', 'Available to MLSS users'],
    'AT-04')

story('US-AT-05', 'View attendance compliance dashboard',
    'Regional SF Officer',
    'a dashboard of attendance compliance rates by school, region, and nationally',
    'I can monitor compliance across my region',
    ['Filterable by school, region, national', 'Visual indicators for compliant vs non-compliant'],
    'AT-05')

story('US-AT-06', 'MLSS read access to attendance data',
    'MLSS Director',
    'read access to PATH student attendance and compliance status',
    'I can view compliance without manual reports',
    ['Read-only access for MLSS users', 'No manual submission required'],
    'AT-06')

story('US-AT-07', 'Flag schools with missing attendance',
    'SF Unit Data Manager',
    'the system to flag schools that have not submitted attendance data on time',
    'I can follow up with non-reporting schools',
    ['Automatic flagging based on expected dates', 'Alert visible on dashboard (DR-06)'],
    'AT-07, DR-06')

doc.add_page_break()

# ===== 4. Meal Management =====
doc.add_heading('4. Meal Management \u2014 Lunch List and Meal Receipt', level=1)

story('US-ML-01', 'Create daily lunch list',
    'SF School Coordinator',
    'create a daily lunch list marking which students will receive a meal (yes/no)',
    'the canteen knows how many meals to prepare',
    ['Auto-filter by programme type', 'Must work offline (NF-09)', 'Count cannot exceed students present (VR-03)', 'Checkbox-based bulk entry'],
    'ML-01, VR-03, NF-09')

story('US-ML-02', 'Record menu selection per student',
    'SF School Coordinator',
    'record each student menu selection where multiple options are offered',
    'meal preferences are tracked',
    ['Multi-select from school menu', 'Only where school offers multiple options'],
    'ML-02')

story('US-ML-03', 'Capture non-PATH meal payment',
    'SF School Coordinator',
    'record payment amount per day for non-PATH/Welfare/Wards students who pay',
    'co-pay revenue is tracked',
    ['Numeric field per student per day', 'Only for non-programme students'],
    'ML-03')

story('US-ML-04', 'View meal turnout statistics',
    'SF Unit Director',
    'see daily, weekly, termly meal turnout (signed up vs. received, by programme type)',
    'I can assess programme reach',
    ['Breakdown by PATH, Welfare, Wards, Other', 'Daily, weekly, termly aggregation', 'School, regional, national level'],
    'ML-04')

story('US-ML-05', 'Track meals per PATH student per week',
    'SF School Coordinator',
    'the system to calculate meals per PATH student per week',
    'I can ensure students eat on all available days',
    ['Per-student weekly meal count', 'Compared against available school days'],
    'ML-05')

story('US-ML-06', 'Flag PATH students not receiving meals',
    'SF Unit Data Manager',
    'the system to flag PATH students consistently not signing up for or receiving meals',
    'follow-up action can be taken',
    ['Configurable threshold', 'Alert on dashboard (DR-06)'],
    'ML-06, DR-06')

doc.add_page_break()

# ===== 5. Canteen =====
doc.add_heading('5. Canteen and Tuck Shop Registry and Validation', level=1)

story('US-CT-01', 'Enter school menu',
    'Canteen Supervisor / SF Coordinator',
    'enter the school menu from a standardized enumeration list',
    'menus are consistent and reportable',
    ['Standardized enumeration list', 'Updated termly'],
    'CT-01')

story('US-CT-02', 'Record menu week of use',
    'Canteen Supervisor',
    'record which weeks each menu is used within a term',
    'menu rotation can be tracked',
    ['Week-of-year selection per menu item', 'Updated termly'],
    'CT-02')

story('US-CT-03', 'Confirm meal receipt per student',
    'Canteen Supervisor',
    'confirm whether each student on the lunch list actually received their meal (yes/no)',
    'there is canteen-level validation of meal distribution',
    ['Cannot exceed meals on lunch list (VR-04)', 'Must work offline (NF-09)', 'Fully functional on tablets (NF-18)', 'Distinct from lunch list sign-up'],
    'CT-03, VR-04, NF-09, NF-18')

story('US-CT-04', 'Record estimated meal cost',
    'Canteen Supervisor / SF Coordinator',
    'record the estimated cost per menu item',
    'meal costing data is available for planning',
    ['Numeric field per menu item', 'Updated termly'],
    'CT-04')

story('US-CT-05', 'Record canteen metadata',
    'School Administrator',
    'record canteen/tuck shop metadata (type, capacity, conditions)',
    'the SF Unit has a profile of canteen facilities',
    ['Fields TBD based on school visit form', 'Updated as needed'],
    'CT-05')

doc.add_page_break()

# ===== 6. ME260 =====
doc.add_heading('6. ME260 Financial Returns', level=1)

story('US-ME-01', 'Complete digital ME260 form',
    'SF School Coordinator / School Bursar',
    'complete the ME260 Nutritional Subsidy Termly Report digitally',
    'I no longer submit paper forms',
    ['Mirrors all data points of paper form (Appendix B)', 'Multi-section: School Info, Meals, Financial Summary, Expenditure, Documents, Signatures'],
    'ME-01')

story('US-ME-02', 'Pre-populated school metadata',
    'SF School Coordinator',
    'the ME260 form pre-populated with school metadata',
    'I avoid re-entering known information',
    ['Auto-populated from school profile and registry', 'Reduces manual entry errors'],
    'ME-02')

story('US-ME-03', 'Enter meals provided and financial summary',
    'SF School Coordinator',
    'enter days served per grade, children served, co-pay, contributions, balances, expenditure',
    'the full financial picture is captured',
    ['Closing balance = brought forward + subsidy + contributions - expenditure (VR-01)', 'Children served cannot exceed enrollment (VR-02)', 'System auto-calculates and validates'],
    'ME-03, VR-01, VR-02')

story('US-ME-04', 'Enter Statement of Expenditure',
    'SF School Coordinator / School Bursar',
    'enter line-item expenditures (date, source, receipt number, amount)',
    'all spending is accounted for',
    ['Add/remove line items', 'Amounts must be positive (VR-07)', 'Auto-calculated total'],
    'ME-04, VR-07')

story('US-ME-05', 'Upload receipts and invoices',
    'SF School Coordinator',
    'upload receipts and invoices (PNG, JPG, PDF)',
    'supporting documentation is attached',
    ['PNG, JPG, PDF supported', 'Camera integration on mobile/tablet', 'Multiple file upload'],
    'ME-05')

story('US-ME-06', 'Apply digital signatures',
    'School Principal / Chairman',
    'apply a digital signature to the ME260 form',
    'the submission is authorized without physical signatures',
    ['Signatures: SF Coordinator, Principal, Chairman', 'Each includes date', 'Replaces physical signatures'],
    'ME-06')

story('US-ME-07', 'ME260 submission and review workflow',
    'SF Unit Reconciliation Officer',
    'a workflow: school submits, I review, I mark Accepted or Requires Resubmission with comments',
    'the review cycle is managed digitally',
    ['School submits form', 'Officer reviews', 'Mark Accepted or Requires Resubmission with comments', 'School notified; can resubmit if needed', 'Acceptance recorded and school notified'],
    'ME-07')

story('US-ME-08', 'Track ME260 compliance status',
    'SF Unit Director',
    'see each school ME260 status per term (Submitted, Under Review, Accepted, Requires Resubmission, Overdue)',
    'I can monitor financial compliance nationally',
    ['Five status values', 'Per-school, per-term tracking'],
    'ME-08')

story('US-ME-09', 'Auto-flag overdue ME260',
    'SF Unit Reconciliation Officer',
    'the system to flag schools that missed the ME260 deadline',
    'I can follow up promptly',
    ['Deadline: second month of following term', 'Alert on dashboard and reconciliation overview'],
    'ME-09')

story('US-ME-10', 'View reconciliation overview',
    'SF Unit Reconciliation Officer',
    'a reconciliation overview of compliance status across all schools by region and term',
    'I have a single view replacing Google Sheets',
    ['Filterable by region and term', 'Shows all statuses', 'Replaces Google Sheets tracker'],
    'ME-10')

story('US-ME-11', 'Track fund balances per school',
    'SF Unit Reconciliation Officer',
    'track fund balances per school (sent vs. disbursed)',
    'I have a single view replacing Excel tracker',
    ['Per-school: subsidy received vs expenditure', 'Replaces Excel tracker'],
    'ME-11')

doc.add_page_break()

# ===== 7. School Visit =====
doc.add_heading('7. School Visit Evaluations', level=1)

story('US-SV-01', 'Pre-populated school visit form',
    'Regional SF Officer',
    'a digital school visit form pre-populated with school profile data',
    'I avoid re-entering known information',
    ['Pre-populated: name, code, address, parish, region, principal, SF coordinator', 'Fully functional on tablets', 'Touch-friendly controls'],
    'SV-01, NF-18')

story('US-SV-02', 'Complete assessment sections',
    'Regional SF Officer',
    'record visit date, financial management, meal preparation, and canteen quality assessments',
    'each visit is consistently evaluated',
    ['Visit date', 'Financial management (TBD)', 'Meal preparation (enum/multi-select)', 'Canteen quality (enum/multi-select)'],
    'SV-02')

story('US-SV-03', 'Assess cook management',
    'Regional SF Officer',
    'capture cook management assessments',
    'cook certification and performance are evaluated',
    ['Fields TBD', 'Includes certification verification'],
    'SV-03')

story('US-SV-04', 'Capture PATH student satisfaction',
    'Regional SF Officer',
    'capture PATH student satisfaction survey results',
    'student feedback is recorded digitally',
    ['Survey fields TBD', 'Linked to school and visit date'],
    'SV-04')

story('US-SV-05', 'Record food source information',
    'Regional SF Officer',
    'record food sources by percentage type',
    'sourcing patterns are tracked',
    ['Percentage allocation by source type', 'Fields TBD'],
    'SV-05')

story('US-SV-06', 'Record food wastage',
    'Regional SF Officer',
    'capture food wastage data',
    'waste levels are monitored',
    ['Enumeration/multi-select', 'Reported termly'],
    'SV-06')

story('US-SV-07', 'Upload visit documentation',
    'Regional SF Officer',
    'upload photos and documents during my visit',
    'visual evidence supports the evaluation',
    ['Camera integration on tablet/mobile', 'Multiple file upload'],
    'SV-07')

story('US-SV-08', 'Add summary feedback',
    'Regional SF Officer',
    'write a summary assessment text',
    'my overall findings are captured',
    ['Free-text field', 'No character limit'],
    'SV-08')

story('US-SV-09', 'Submit evaluation',
    'Regional SF Officer',
    'submit the completed evaluation',
    'it is visible to HQ SF Unit and the assessed school',
    ['Visible to HQ SF Unit', 'Visible to assessed school', 'Confirmation prompt before submission'],
    'SV-09')

story('US-SV-10', 'School visit summary reports',
    'SF Unit Director',
    'summary reports from visits aggregated by region and nationally',
    'I can assess programme quality at scale',
    ['Aggregated by region and nationally', 'Exportable in PDF and Excel/CSV'],
    'SV-10')

doc.add_page_break()

# ===== 8. Dashboards =====
doc.add_heading('8. Dashboards and Reporting', level=1)

story('US-DR-01', 'Role-based dashboards',
    'System Administrator',
    'role-based dashboards at school, regional, and national levels',
    'each user sees data relevant to their scope',
    ['School users see own school', 'Regional users see their region', 'HQ users see national with drill-down'],
    'DR-01')

story('US-DR-02', 'SF indicators on dashboard',
    'SF Unit Director',
    'key SF indicators on dashboards (attendance, children served, feeding days, co-pay, meal size, menu, cost, losses, wastage, sourcing)',
    'I can monitor programme performance',
    ['All indicators from BRD Section 12.1', 'Frequency as specified', 'Charts and tables'],
    'DR-02, Section 12.1')

story('US-DR-03', 'Nutrition indicators on dashboard',
    'Regional SF Officer',
    'nutrition indicators (overweight/underweight, nutritional requirements, policy compliance) on dashboards',
    'nutritional status is visible',
    ['Aggregated by school and region', 'Policy compliance indicators'],
    'DR-03')

story('US-DR-04', 'Filter dashboards',
    'PIOJ Director',
    'filter dashboards by school, parish, region, term, and school year',
    'I can drill into specific areas',
    ['All five filter dimensions', 'Filters combinable'],
    'DR-04')

story('US-DR-05', 'Export reports',
    'SF Unit Data Manager',
    'export reports in PDF and Excel/CSV',
    'I can share data with stakeholders without system access',
    ['PDF export', 'Excel/CSV export', 'All report types'],
    'DR-05')

story('US-DR-06', 'Alerts and notifications',
    'SF Unit Director',
    'alerts for overdue ME260, missing attendance, low PATH attendance, low meal turnout',
    'I am proactively informed of issues',
    ['Overdue ME260 alert', 'Missing attendance alert', 'PATH non-compliance alert', 'Low meal turnout alert', 'Visible on dashboard'],
    'DR-06')

doc.add_page_break()

# ===== 9. Cross-Cutting =====
doc.add_heading('9. Cross-Cutting / Non-Functional User Stories', level=1)

story('US-RBAC-01', 'Role-based access control',
    'System Administrator (ESMS Coordinator)',
    'assign roles with configurable CRUD permissions per data domain, scoped by level (National, Regional, School)',
    'users only access authorized data',
    ['One or more roles per user', 'CRUD permissions per data domain', 'National: all schools/regions', 'Regional: assigned region only', 'School: own school only', 'MLSS: read attendance/compliance, write PATH lists', 'Audit logging of all modifications'],
    'Section 3.3')

story('US-NF-01', 'Offline data entry',
    'SF School Coordinator',
    'enter attendance, lunch lists, and meal confirmations offline with auto-sync',
    'schools with poor internet can use the system daily',
    ['Offline: attendance, lunch list, meal confirmation', 'Auto-sync on reconnect', 'Secure local storage', 'Conflict resolution'],
    'NF-08, NF-09, NF-10')

story('US-NF-02', 'Audit trail',
    'SF Unit Data Manager',
    'a complete audit trail of all data modifications (user, timestamp, previous values)',
    'any change can be traced',
    ['Every modification logged', 'User, timestamp, old/new values', 'Accessible to authorized users'],
    'NF-15, Section 3.3')

story('US-NF-03', 'Mobile-responsive interface',
    'Regional SF Officer',
    'the system responsive on tablets and smartphones with touch-friendly controls',
    'I can complete evaluations in the field',
    ['Responsive for tablets and smartphones', 'Touch-friendly (larger buttons, swipe)', 'Camera integration for file upload', 'Chrome, Firefox, Edge, Safari support'],
    'NF-17, NF-18, Section 11.3')

story('US-INT-01', 'Bulk PATH import with validation',
    'ESMS Coordinator',
    'import PATH lists via CSV/Excel with validation and error reporting',
    'data quality is maintained during batch updates',
    ['CSV and Excel supported', 'Validation: invalid IDs, mismatched codes', 'Error report after import', 'Logged with timestamp, source, user', 'RESTful API for future real-time exchange'],
    'Section 9.2, 9.3')

doc.save('c:/DomMIS/SF_eMIS_User_Stories.docx')
print('Done - saved to c:/DomMIS/SF_eMIS_User_Stories.docx')
